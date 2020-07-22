from docx import Document
from docx.section import Section
from docx.shared import Pt, Inches
from random import random
import os

if not os.path.exists('./out'):
    os.makedirs('./out')

# Generate text

doc = Document()
doc.styles['Normal'].font.name = 'Courier New'
doc.styles['Normal'].font.size = Pt(9)
doc.styles['Normal'].paragraph_format.space_before = Pt(0)
doc.styles['Normal'].paragraph_format.space_after = Pt(0)
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(6.25)
    section.page_height = Inches(9.25)

letters = [l for l in "abcdefghijklmnopqrstuvwxyz .,"]
length = len(letters)
width = 56

odd = True
for i in range(410):
    num = str(i+1)
    print('Page ' + num)
    for j in range(46):
        line = ''.join([letters[int(random() * length)] for _ in range(width)])
        doc.add_paragraph(line)
    doc.add_paragraph(' ')
    doc.add_paragraph(' ')
    if odd is True:
        doc.add_paragraph(num)
    else:
        doc.add_paragraph((' ' * (width-len(num))) + num)
    odd = not odd

doc.save('out/book.docx')

# Generate hard cover

from PIL import Image, ImageFont, ImageDraw

size = (2956, 2100)
background_color = (27, 45, 8)
text_color = (230, 230, 230)
title_font = ImageFont.truetype('cour.ttf', 84)
spine_font = ImageFont.truetype('courbd.ttf', 60)
text = ''.join([letters[int(random() * length)] for _ in range(16)])

img = Image.new('RGB', size, color=background_color)
draw = ImageDraw.Draw(img)
draw.text((200, 200), text, fill=text_color, font=title_font)

img.save(r'out/cover.pdf')